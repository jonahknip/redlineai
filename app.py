"""
PlanSet Review Web App
Simple web interface for civil engineering planset PDF review
"""

import os
import tempfile
import logging
import re
import traceback
import json
import uuid
from datetime import datetime
from pathlib import Path
from io import BytesIO

from flask import Flask, render_template, request, jsonify, send_file, session
from werkzeug.utils import secure_filename

from agent.plan_reviewer import CivilEngineeringPMAgent

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

app = Flask(__name__)
app.config['MAX_CONTENT_LENGTH'] = 500 * 1024 * 1024  # 500MB max upload
app.secret_key = os.environ.get('SECRET_KEY', 'planset-review-secret-key-change-in-prod')

# In-memory storage for review history (per session)
review_history = {}

# Allowed extensions
ALLOWED_EXTENSIONS = {'.pdf'}

# OneDrive/SharePoint URL patterns
SHARE_URL_PATTERNS = [
    r'https?://[a-zA-Z0-9-]+\.sharepoint\.com/.*',
    r'https?://[a-zA-Z0-9-]+-my\.sharepoint\.com/.*',
    r'https?://onedrive\.live\.com/.*',
    r'https?://1drv\.ms/.*',
]


def allowed_file(filename: str) -> bool:
    """Check if file has allowed extension"""
    return Path(filename).suffix.lower() in ALLOWED_EXTENSIONS


def is_share_url(text: str) -> bool:
    """Check if text is a OneDrive/SharePoint URL"""
    for pattern in SHARE_URL_PATTERNS:
        if re.match(pattern, text.strip(), re.IGNORECASE):
            return True
    return False


def download_from_share_url(url: str) -> tuple[str, str]:
    """
    Download a file from OneDrive/SharePoint sharing URL
    
    Returns tuple of (local_path, filename)
    """
    import requests
    
    url = url.strip()
    logger.info(f"Attempting to download from: {url}")
    
    # Try to get a direct download URL
    # For OneDrive/SharePoint, we can often modify the URL to get direct download
    download_url = url
    
    # Handle different URL formats
    if '1drv.ms' in url:
        # Short URL - need to follow redirect
        response = requests.head(url, allow_redirects=True)
        download_url = response.url
    
    # Try to convert to direct download URL
    if 'sharepoint.com' in download_url or 'onedrive.live.com' in download_url:
        # Replace sharing indicator with download
        if '?e=' in download_url:
            download_url = download_url.split('?e=')[0]
        if ':b:' in download_url:
            # This is a file link, try to get download URL
            download_url = download_url.replace(':b:', ':b:/') + '?download=1'
        elif 'download=1' not in download_url:
            separator = '&' if '?' in download_url else '?'
            download_url = download_url + separator + 'download=1'
    
    logger.info(f"Download URL: {download_url}")
    
    # Download the file
    headers = {
        'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36'
    }
    
    response = requests.get(download_url, headers=headers, stream=True, timeout=300)
    
    if response.status_code != 200:
        raise Exception(f"Failed to download file: HTTP {response.status_code}")
    
    # Try to get filename from headers
    content_disposition = response.headers.get('Content-Disposition', '')
    filename = 'planset.pdf'
    
    if 'filename=' in content_disposition:
        # Extract filename from header
        match = re.search(r'filename[*]?=["\']?([^"\';\n]+)', content_disposition)
        if match:
            filename = match.group(1).strip()
    
    # Verify it's a PDF
    content_type = response.headers.get('Content-Type', '')
    if 'pdf' not in content_type.lower() and not filename.lower().endswith('.pdf'):
        # Check first bytes for PDF signature
        first_bytes = response.content[:5]
        if first_bytes != b'%PDF-':
            raise Exception("The shared file does not appear to be a PDF")
    
    # Save to temp file
    temp_dir = tempfile.mkdtemp()
    safe_filename = secure_filename(filename)
    if not safe_filename.lower().endswith('.pdf'):
        safe_filename += '.pdf'
    
    local_path = os.path.join(temp_dir, safe_filename)
    
    with open(local_path, 'wb') as f:
        f.write(response.content)
    
    file_size = os.path.getsize(local_path)
    logger.info(f"Downloaded {safe_filename} ({file_size} bytes)")
    
    return local_path, safe_filename


def analyze_planset(pdf_path: str, use_vision: bool = True, checklist: dict = None, custom_instructions: str = "") -> dict:
    """
    Analyze a planset PDF and return results
    """
    try:
        agent = CivilEngineeringPMAgent(pdf_path)
        
        # Get page count
        page_count = len(agent.doc)
        
        # Generate AI-powered HTML report (falls back to basic if no API key)
        report = agent.generate_ai_report(
            use_vision=use_vision,
            checklist=checklist,
            custom_instructions=custom_instructions
        )
        
        # Check if report is HTML or plain text
        is_html = report.strip().startswith('<')
        
        # Get JSON data for structured display
        json_data = agent.export_json()
        
        # Extract project name for history
        project_name = json_data.get('project_info', {}).get('project_name', 'Unknown Project')
        
        return {
            'success': True,
            'page_count': page_count,
            'report': report,
            'is_html': is_html,
            'data': json_data,
            'project_name': project_name
        }
    except Exception as e:
        logger.error(f"Analysis error: {e}\n{traceback.format_exc()}")
        return {
            'success': False,
            'error': str(e)
        }


@app.route('/')
def index():
    """Main page"""
    return render_template('index.html')


@app.route('/api/review', methods=['POST'])
def review_planset():
    """
    Review a planset from either file upload or URL
    """
    temp_path = None
    
    try:
        # Check if URL was provided
        share_url = request.form.get('url', '').strip()
        
        if share_url:
            # Download from OneDrive/SharePoint
            if not is_share_url(share_url):
                return jsonify({
                    'success': False,
                    'error': 'Invalid URL. Please provide a OneDrive or SharePoint sharing link.'
                }), 400
            
            try:
                temp_path, filename = download_from_share_url(share_url)
            except Exception as e:
                logger.error(f"Download error: {e}")
                return jsonify({
                    'success': False,
                    'error': f'Failed to download file: {str(e)}'
                }), 400
        
        elif 'file' in request.files:
            # Handle file upload
            file = request.files['file']
            
            if file.filename == '':
                return jsonify({
                    'success': False,
                    'error': 'No file selected'
                }), 400
            
            if not allowed_file(file.filename):
                return jsonify({
                    'success': False,
                    'error': 'Invalid file type. Please upload a PDF file.'
                }), 400
            
            # Save uploaded file
            temp_dir = tempfile.mkdtemp()
            filename = secure_filename(file.filename)
            temp_path = os.path.join(temp_dir, filename)
            file.save(temp_path)
            
            logger.info(f"Saved uploaded file: {filename}")
        
        else:
            return jsonify({
                'success': False,
                'error': 'Please upload a PDF file or provide a OneDrive/SharePoint link.'
            }), 400
        
        # Get optional parameters
        use_vision = request.form.get('use_vision', 'true').lower() == 'true'
        custom_instructions = request.form.get('instructions', '')
        checklist_id = request.form.get('checklist', '')
        
        # Load checklist if specified
        checklist = None
        if checklist_id:
            checklists = get_checklists().get_json().get('checklists', {})
            checklist = checklists.get(checklist_id)
        
        # Analyze the planset
        result = analyze_planset(
            temp_path,
            use_vision=use_vision,
            checklist=checklist,
            custom_instructions=custom_instructions
        )
        
        if result['success']:
            return jsonify(result)
        else:
            return jsonify(result), 500
    
    except Exception as e:
        logger.error(f"Review error: {e}\n{traceback.format_exc()}")
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500
    
    finally:
        # Clean up temp file
        if temp_path and os.path.exists(temp_path):
            try:
                os.remove(temp_path)
                os.rmdir(os.path.dirname(temp_path))
            except Exception:
                pass


@app.route('/health')
def health():
    """Health check endpoint"""
    return jsonify({'status': 'healthy'})


@app.route('/api/export/word', methods=['POST'])
def export_word():
    """Export report as Word document"""
    try:
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        import html
        from bs4 import BeautifulSoup
        
        data = request.get_json()
        report_html = data.get('report', '')
        project_name = data.get('project_name', 'Planset Review')
        
        # Create Word document
        doc = Document()
        
        # Parse HTML and convert to Word
        soup = BeautifulSoup(report_html, 'html.parser')
        
        # Add title
        title = doc.add_heading(f'{project_name} - Review Report', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add date
        date_para = doc.add_paragraph(f'Generated: {datetime.now().strftime("%B %d, %Y")}')
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()  # Spacer
        
        # Process HTML content
        for element in soup.find_all(['h1', 'h2', 'h3', 'p', 'ul', 'ol', 'table', 'div']):
            if element.name in ['h1', 'h2', 'h3']:
                level = int(element.name[1])
                doc.add_heading(element.get_text(strip=True), level)
            elif element.name == 'p':
                doc.add_paragraph(element.get_text(strip=True))
            elif element.name in ['ul', 'ol']:
                for li in element.find_all('li', recursive=False):
                    para = doc.add_paragraph(li.get_text(strip=True), style='List Bullet')
            elif element.name == 'table':
                rows = element.find_all('tr')
                if rows:
                    cols = len(rows[0].find_all(['td', 'th']))
                    if cols > 0:
                        table = doc.add_table(rows=len(rows), cols=cols)
                        table.style = 'Table Grid'
                        for i, row in enumerate(rows):
                            cells = row.find_all(['td', 'th'])
                            for j, cell in enumerate(cells):
                                if j < cols:
                                    table.rows[i].cells[j].text = cell.get_text(strip=True)
        
        # Save to BytesIO
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        filename = f"{project_name.replace(' ', '_')}_Review_{datetime.now().strftime('%Y%m%d')}.docx"
        
        return send_file(
            buffer,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            as_attachment=True,
            download_name=filename
        )
        
    except Exception as e:
        logger.error(f"Word export error: {e}\n{traceback.format_exc()}")
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/api/export/pdf', methods=['POST'])
def export_pdf():
    """Export report as PDF document"""
    try:
        data = request.get_json()
        report_html = data.get('report', '')
        project_name = data.get('project_name', 'Planset Review')
        
        # Create full HTML document with styling
        full_html = f"""
        <!DOCTYPE html>
        <html>
        <head>
            <meta charset="UTF-8">
            <style>
                body {{
                    font-family: 'Segoe UI', Arial, sans-serif;
                    font-size: 11pt;
                    line-height: 1.5;
                    color: #333;
                    max-width: 800px;
                    margin: 0 auto;
                    padding: 40px;
                }}
                h1 {{
                    color: #1B365D;
                    border-bottom: 3px solid #C8102E;
                    padding-bottom: 10px;
                }}
                h2 {{
                    color: #1B365D;
                    border-bottom: 1px solid #ddd;
                    padding-bottom: 5px;
                    margin-top: 25px;
                }}
                h3 {{
                    color: #333;
                    margin-top: 20px;
                }}
                table {{
                    width: 100%;
                    border-collapse: collapse;
                    margin: 15px 0;
                }}
                th, td {{
                    border: 1px solid #ddd;
                    padding: 8px;
                    text-align: left;
                }}
                th {{
                    background: #f5f5f5;
                    color: #1B365D;
                }}
                ul, ol {{
                    margin: 10px 0;
                    padding-left: 25px;
                }}
                .critical, .red {{
                    color: #C8102E;
                    font-weight: bold;
                }}
                .warning, .orange {{
                    color: #E67E22;
                    font-weight: bold;
                }}
                .success, .green {{
                    color: #27AE60;
                    font-weight: bold;
                }}
                .header {{
                    text-align: center;
                    margin-bottom: 30px;
                }}
                .footer {{
                    text-align: center;
                    margin-top: 40px;
                    font-size: 9pt;
                    color: #999;
                }}
            </style>
        </head>
        <body>
            <div class="header">
                <h1>{project_name}</h1>
                <p>Review Report - Generated {datetime.now().strftime("%B %d, %Y")}</p>
            </div>
            {report_html}
            <div class="footer">
                <p>Generated by PlanSet Review Agent - Abonmarche</p>
            </div>
        </body>
        </html>
        """
        
        # Try WeasyPrint first, fall back to basic HTML
        try:
            from weasyprint import HTML
            buffer = BytesIO()
            HTML(string=full_html).write_pdf(buffer)
            buffer.seek(0)
            mimetype = 'application/pdf'
            extension = 'pdf'
        except ImportError:
            # Fallback: return HTML file if WeasyPrint not available
            buffer = BytesIO(full_html.encode('utf-8'))
            mimetype = 'text/html'
            extension = 'html'
        
        filename = f"{project_name.replace(' ', '_')}_Review_{datetime.now().strftime('%Y%m%d')}.{extension}"
        
        return send_file(
            buffer,
            mimetype=mimetype,
            as_attachment=True,
            download_name=filename
        )
        
    except Exception as e:
        logger.error(f"PDF export error: {e}\n{traceback.format_exc()}")
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/api/history', methods=['GET'])
def get_history():
    """Get review history for current session"""
    session_id = session.get('session_id')
    if not session_id:
        session_id = str(uuid.uuid4())
        session['session_id'] = session_id
    
    history = review_history.get(session_id, [])
    return jsonify({'success': True, 'history': history})


@app.route('/api/history', methods=['POST'])
def save_to_history():
    """Save a review to history"""
    try:
        session_id = session.get('session_id')
        if not session_id:
            session_id = str(uuid.uuid4())
            session['session_id'] = session_id
        
        data = request.get_json()
        
        review_entry = {
            'id': str(uuid.uuid4()),
            'timestamp': datetime.now().isoformat(),
            'project_name': data.get('project_name', 'Unknown Project'),
            'page_count': data.get('page_count', 0),
            'report': data.get('report', ''),
            'is_html': data.get('is_html', False),
            'data': data.get('data', {})
        }
        
        if session_id not in review_history:
            review_history[session_id] = []
        
        # Keep last 20 reviews
        review_history[session_id].insert(0, review_entry)
        review_history[session_id] = review_history[session_id][:20]
        
        return jsonify({'success': True, 'id': review_entry['id']})
        
    except Exception as e:
        logger.error(f"Save history error: {e}")
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/api/history/<review_id>', methods=['GET'])
def get_review(review_id):
    """Get a specific review from history"""
    session_id = session.get('session_id')
    if not session_id:
        return jsonify({'success': False, 'error': 'No session'}), 404
    
    history = review_history.get(session_id, [])
    for review in history:
        if review['id'] == review_id:
            return jsonify({'success': True, 'review': review})
    
    return jsonify({'success': False, 'error': 'Review not found'}), 404


@app.route('/api/email', methods=['POST'])
def send_email():
    """Send report via email"""
    try:
        data = request.get_json()
        recipient = data.get('email', '')
        report_html = data.get('report', '')
        project_name = data.get('project_name', 'Planset Review')
        
        if not recipient:
            return jsonify({'success': False, 'error': 'Email address required'}), 400
        
        # For now, we'll use a simple mailto link approach
        # In production, you'd configure SMTP settings
        import urllib.parse
        
        # Create plain text version
        from bs4 import BeautifulSoup
        soup = BeautifulSoup(report_html, 'html.parser')
        plain_text = soup.get_text(separator='\n')
        
        subject = f"PlanSet Review Report: {project_name}"
        body = f"Please find the review report for {project_name} below:\n\n{plain_text[:2000]}..."
        
        mailto_link = f"mailto:{recipient}?subject={urllib.parse.quote(subject)}&body={urllib.parse.quote(body)}"
        
        return jsonify({
            'success': True, 
            'mailto_link': mailto_link,
            'message': 'Email link generated. Click to open your email client.'
        })
        
    except Exception as e:
        logger.error(f"Email error: {e}")
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/api/checklist', methods=['GET'])
def get_checklists():
    """Get available checklist templates"""
    checklists = {
        'standard': {
            'name': 'Standard Civil Review',
            'items': [
                {'id': 'cover', 'text': 'Cover sheet with project info', 'required': True},
                {'id': 'index', 'text': 'Sheet index present', 'required': True},
                {'id': 'pe_seal', 'text': 'PE seal and signature', 'required': True},
                {'id': 'scale', 'text': 'Scale indicated on all sheets', 'required': True},
                {'id': 'north', 'text': 'North arrow on plan sheets', 'required': True},
                {'id': 'legend', 'text': 'Legend/symbols defined', 'required': False},
                {'id': 'grading', 'text': 'Grading plan included', 'required': True},
                {'id': 'drainage', 'text': 'Drainage plan included', 'required': True},
                {'id': 'erosion', 'text': 'Erosion control plan', 'required': True},
                {'id': 'utilities', 'text': 'Utility plan included', 'required': False},
                {'id': 'details', 'text': 'Construction details', 'required': True},
                {'id': 'specs', 'text': 'Reference to specifications', 'required': False},
            ]
        },
        'mdot': {
            'name': 'MDOT Project Review',
            'items': [
                {'id': 'cover', 'text': 'MDOT standard cover sheet', 'required': True},
                {'id': 'index', 'text': 'Sheet index per MDOT format', 'required': True},
                {'id': 'pe_seal', 'text': 'Michigan PE seal', 'required': True},
                {'id': 'typical', 'text': 'Typical sections', 'required': True},
                {'id': 'mot', 'text': 'Maintenance of Traffic plan', 'required': True},
                {'id': 'signing', 'text': 'Signing and striping plan', 'required': True},
                {'id': 'drainage', 'text': 'Drainage plan with calculations', 'required': True},
                {'id': 'soil_boring', 'text': 'Soil boring locations', 'required': True},
                {'id': 'row', 'text': 'ROW lines shown', 'required': True},
                {'id': 'permits', 'text': 'Permit requirements noted', 'required': True},
                {'id': 'quantities', 'text': 'Pay item quantities', 'required': True},
                {'id': 'cross_sections', 'text': 'Cross sections provided', 'required': True},
            ]
        },
        'municipal': {
            'name': 'Municipal Infrastructure',
            'items': [
                {'id': 'cover', 'text': 'Cover sheet with municipal info', 'required': True},
                {'id': 'survey', 'text': 'Survey/existing conditions', 'required': True},
                {'id': 'layout', 'text': 'Site layout plan', 'required': True},
                {'id': 'grading', 'text': 'Grading and drainage', 'required': True},
                {'id': 'water', 'text': 'Water main plan', 'required': True},
                {'id': 'sanitary', 'text': 'Sanitary sewer plan', 'required': True},
                {'id': 'storm', 'text': 'Storm sewer plan', 'required': True},
                {'id': 'paving', 'text': 'Paving plan', 'required': True},
                {'id': 'landscape', 'text': 'Landscape plan', 'required': False},
                {'id': 'lighting', 'text': 'Street lighting plan', 'required': False},
                {'id': 'swppp', 'text': 'SWPPP/erosion control', 'required': True},
                {'id': 'details', 'text': 'Standard details', 'required': True},
            ]
        }
    }
    return jsonify({'success': True, 'checklists': checklists})


if __name__ == '__main__':
    port = int(os.environ.get('PORT', 5000))
    debug = os.environ.get('DEBUG', 'false').lower() == 'true'
    app.run(host='0.0.0.0', port=port, debug=debug)
