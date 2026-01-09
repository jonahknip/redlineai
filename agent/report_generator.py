"""
Report Generator - Generate completed checklist reports matching Abonmarche form format
"""

import io
import os
from typing import Dict, Any, List, Optional
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


class ReportGenerator:
    """
    Generates completed QA/QC checklist reports matching exact Abonmarche format.
    """
    
    def __init__(self):
        """Initialize the report generator."""
        pass
    
    def generate_html_report(self, review_results: Dict[str, Any]) -> str:
        """
        Generate an HTML report that exactly matches the Abonmarche fillable form style.
        """
        project = review_results.get('project_summary', {})
        summary = review_results.get('summary', {})
        phase = review_results.get('phase', '')
        
        # Get form metadata
        form_project_name = review_results.get('form_project_name') or project.get('project_name', '')
        form_project_number = review_results.get('form_project_number') or project.get('project_number', '')
        form_project_manager = review_results.get('form_project_manager', '')
        form_reviewer = review_results.get('form_reviewer', 'RedlineAI')
        form_date = datetime.now().strftime('%m/%d/%Y')
        
        html = f'''<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{phase} Engineering QA/QC Review - {form_project_name}</title>
    <style>
        @import url('https://fonts.googleapis.com/css2?family=Inter:wght@400;500;600;700&display=swap');
        
        * {{
            box-sizing: border-box;
            margin: 0;
            padding: 0;
        }}
        
        body {{
            font-family: 'Inter', Arial, sans-serif;
            font-size: 11px;
            line-height: 1.4;
            color: #1a1a1a;
            background: #f5f5f5;
            padding: 20px;
        }}
        
        .report-container {{
            max-width: 850px;
            margin: 0 auto;
            background: white;
            box-shadow: 0 2px 20px rgba(0,0,0,0.1);
        }}
        
        .page {{
            padding: 40px 50px;
            min-height: 1100px;
            position: relative;
        }}
        
        /* Header */
        .header {{
            display: flex;
            justify-content: space-between;
            align-items: flex-start;
            margin-bottom: 8px;
            padding-bottom: 12px;
            border-bottom: 2px solid #C8102E;
        }}
        
        .logo {{
            display: flex;
            align-items: center;
            gap: 8px;
        }}
        
        .logo-mark {{
            display: flex;
            gap: 3px;
        }}
        
        .logo-mark .bar {{
            width: 12px;
            height: 36px;
            background: #C8102E;
        }}
        
        .logo-text {{
            font-size: 28px;
            font-weight: 700;
            color: #1a365d;
            letter-spacing: -0.5px;
        }}
        
        .header-tagline {{
            font-size: 11px;
            color: #666;
            text-align: right;
        }}
        
        /* Title */
        .form-title {{
            font-size: 22px;
            font-weight: 700;
            color: #1a1a1a;
            margin: 30px 0 25px;
        }}
        
        /* Project Info */
        .project-info {{
            margin-bottom: 30px;
        }}
        
        .info-row {{
            display: flex;
            align-items: baseline;
            margin-bottom: 8px;
        }}
        
        .info-label {{
            font-weight: 500;
            min-width: 120px;
            color: #1a1a1a;
        }}
        
        .info-value {{
            flex: 1;
            border-bottom: 1px solid #333;
            padding-bottom: 2px;
            min-height: 18px;
        }}
        
        /* Section */
        .section {{
            margin-bottom: 25px;
        }}
        
        .section-title {{
            font-size: 12px;
            font-weight: 700;
            color: #1a1a1a;
            text-transform: uppercase;
            margin-bottom: 10px;
            padding-bottom: 4px;
            border-bottom: 1px solid #1a1a1a;
        }}
        
        .section-note {{
            font-size: 10px;
            font-style: italic;
            color: #666;
            margin-bottom: 10px;
        }}
        
        /* Checklist Table */
        .checklist-table {{
            width: 100%;
            border-collapse: collapse;
            font-size: 10px;
        }}
        
        .checklist-table th {{
            text-align: center;
            font-weight: 600;
            padding: 8px 4px;
            border-bottom: 1px solid #333;
            text-decoration: underline;
        }}
        
        .checklist-table th.item-col {{
            text-align: left;
            text-decoration: none;
            width: 55%;
        }}
        
        .checklist-table th.check-col {{
            width: 8%;
        }}
        
        .checklist-table th.comments-col {{
            text-align: left;
            width: 21%;
        }}
        
        .checklist-table td {{
            padding: 6px 4px;
            vertical-align: top;
            border-bottom: 1px solid #e0e0e0;
        }}
        
        .checklist-table td.item-text {{
            padding-right: 10px;
        }}
        
        .checklist-table td.check-cell {{
            text-align: center;
        }}
        
        .checkbox {{
            width: 14px;
            height: 14px;
            border: 1px solid #333;
            display: inline-flex;
            align-items: center;
            justify-content: center;
            font-size: 11px;
            font-weight: 700;
        }}
        
        .checkbox.checked {{
            background: #e8f5e9;
            color: #2e7d32;
        }}
        
        .checkbox.checked-no {{
            background: #ffebee;
            color: #c62828;
        }}
        
        .checkbox.checked-na {{
            background: #f5f5f5;
            color: #666;
        }}
        
        .comments-cell {{
            font-size: 9px;
            color: #444;
            max-width: 180px;
            word-wrap: break-word;
        }}
        
        /* Review Comments Section */
        .review-comments {{
            margin-top: 40px;
            padding-top: 20px;
        }}
        
        .review-comments-title {{
            font-size: 12px;
            font-weight: 700;
            text-decoration: underline;
            margin-bottom: 15px;
        }}
        
        .review-comments-content {{
            min-height: 150px;
            border-bottom: 1px solid #333;
            padding: 10px 0;
        }}
        
        .comment-item {{
            margin-bottom: 10px;
            padding: 8px 12px;
            background: #fff8e1;
            border-left: 3px solid #ffc107;
            font-size: 10px;
        }}
        
        .comment-item.critical {{
            background: #ffebee;
            border-left-color: #c62828;
        }}
        
        .comment-item strong {{
            color: #c62828;
        }}
        
        /* Footer */
        .footer {{
            position: absolute;
            bottom: 30px;
            left: 50px;
            right: 50px;
        }}
        
        .footer-logo {{
            display: flex;
            gap: 3px;
        }}
        
        .footer-logo .bar {{
            width: 8px;
            height: 24px;
            background: #C8102E;
        }}
        
        .footer-text {{
            text-align: center;
            font-size: 10px;
            color: #666;
            margin-top: 15px;
        }}
        
        /* Summary Box */
        .summary-box {{
            display: flex;
            gap: 20px;
            margin: 20px 0;
            padding: 15px;
            background: #f8f9fa;
            border-radius: 6px;
        }}
        
        .summary-stat {{
            text-align: center;
            flex: 1;
        }}
        
        .summary-stat .number {{
            font-size: 28px;
            font-weight: 700;
        }}
        
        .summary-stat .label {{
            font-size: 10px;
            text-transform: uppercase;
            color: #666;
        }}
        
        .summary-stat.total .number {{ color: #1a365d; }}
        .summary-stat.passed .number {{ color: #2e7d32; }}
        .summary-stat.failed .number {{ color: #c62828; }}
        .summary-stat.na .number {{ color: #666; }}
        
        /* Items Required Section */
        .items-required {{
            margin-top: 30px;
            padding: 20px;
            background: #fff8e1;
            border: 2px solid #ffc107;
            border-radius: 6px;
        }}
        
        .items-required h3 {{
            color: #f57f17;
            font-size: 14px;
            margin-bottom: 15px;
        }}
        
        .items-required ul {{
            list-style: none;
        }}
        
        .items-required li {{
            padding: 8px 0;
            border-bottom: 1px solid #ffe082;
            font-size: 11px;
        }}
        
        .items-required li:last-child {{
            border-bottom: none;
        }}
        
        .items-required .item-id {{
            font-weight: 700;
            color: #c62828;
        }}
        
        /* Print styles */
        @media print {{
            body {{
                background: white;
                padding: 0;
            }}
            
            .report-container {{
                box-shadow: none;
            }}
            
            .page {{
                padding: 20px 30px;
                min-height: auto;
                page-break-after: always;
            }}
            
            .summary-box {{
                -webkit-print-color-adjust: exact;
                print-color-adjust: exact;
            }}
        }}
    </style>
</head>
<body>
    <div class="report-container">
        <div class="page">
            <!-- Header -->
            <div class="header">
                <div class="logo">
                    <div class="logo-mark">
                        <div class="bar"></div>
                        <div class="bar"></div>
                    </div>
                    <div class="logo-text">ABONMARCHE</div>
                </div>
                <div class="header-tagline">
                    Engineering &bull; Architecture &bull; Land Surveying
                </div>
            </div>
            
            <!-- Form Title -->
            <h1 class="form-title">{phase} ENGINEERING QA/QC Review</h1>
            
            <!-- Project Info -->
            <div class="project-info">
                <div class="info-row">
                    <span class="info-label">Project No.:</span>
                    <span class="info-value">{form_project_number}</span>
                </div>
                <div class="info-row">
                    <span class="info-label">Project Name:</span>
                    <span class="info-value">{form_project_name}</span>
                </div>
                <div class="info-row">
                    <span class="info-label">Project Manager:</span>
                    <span class="info-value">{form_project_manager}</span>
                </div>
                <div class="info-row">
                    <span class="info-label">Reviewer:</span>
                    <span class="info-value">{form_reviewer}</span>
                </div>
                <div class="info-row">
                    <span class="info-label">Date:</span>
                    <span class="info-value">{form_date}</span>
                </div>
            </div>
            
            <!-- Summary -->
            <div class="summary-box">
                <div class="summary-stat total">
                    <div class="number">{summary.get('total_items', 0)}</div>
                    <div class="label">Total Items</div>
                </div>
                <div class="summary-stat passed">
                    <div class="number">{summary.get('yes_count', 0)}</div>
                    <div class="label">Passed</div>
                </div>
                <div class="summary-stat failed">
                    <div class="number">{summary.get('no_count', 0)}</div>
                    <div class="label">Failed</div>
                </div>
                <div class="summary-stat na">
                    <div class="number">{summary.get('na_count', 0)}</div>
                    <div class="label">N/A</div>
                </div>
            </div>
'''
        
        # Add each section
        for section in review_results.get('sections', []):
            html += f'''
            <div class="section">
                <div class="section-title">{section['title']}</div>
                <table class="checklist-table">
                    <thead>
                        <tr>
                            <th class="item-col"></th>
                            <th class="check-col">YES</th>
                            <th class="check-col">NO</th>
                            <th class="check-col">N/A</th>
                            <th class="comments-col">COMMENTS</th>
                        </tr>
                    </thead>
                    <tbody>
'''
            for item in section.get('items', []):
                status = item.get('status', 'N/A')
                comments = item.get('comments', '')
                if len(comments) > 100:
                    comments = comments[:100] + '...'
                
                yes_check = '<span class="checkbox checked">&#10004;</span>' if status == 'YES' else '<span class="checkbox"></span>'
                no_check = '<span class="checkbox checked-no">&#10004;</span>' if status == 'NO' else '<span class="checkbox"></span>'
                na_check = '<span class="checkbox checked-na">&#10004;</span>' if status == 'N/A' else '<span class="checkbox"></span>'
                
                html += f'''
                        <tr>
                            <td class="item-text">{item.get('text', '')}</td>
                            <td class="check-cell">{yes_check}</td>
                            <td class="check-cell">{no_check}</td>
                            <td class="check-cell">{na_check}</td>
                            <td class="comments-cell">{comments}</td>
                        </tr>
'''
            
            html += '''
                    </tbody>
                </table>
            </div>
'''
        
        # Items required for next phase
        items_for_next = summary.get('items_for_next_phase', [])
        next_phase = review_results.get('next_phase')
        
        if items_for_next and next_phase:
            html += f'''
            <div class="items-required">
                <h3>Items Required to Advance to {next_phase} Phase ({len(items_for_next)} items)</h3>
                <ul>
'''
            for item in items_for_next:
                html += f'''
                    <li>
                        <span class="item-id">{item.get('id', '')}:</span> {item.get('text', '')}
                    </li>
'''
            html += '''
                </ul>
            </div>
'''
        
        # Review Comments Section
        html += '''
            <div class="review-comments">
                <div class="review-comments-title">Review Comments</div>
                <div class="review-comments-content">
'''
        
        # Add failed items as comments
        for section in review_results.get('sections', []):
            for item in section.get('items', []):
                if item.get('status') == 'NO':
                    comments = item.get('comments', 'Issue identified')
                    if len(comments) > 200:
                        comments = comments[:200] + '...'
                    html += f'''
                    <div class="comment-item critical">
                        <strong>{item.get('id', '')}:</strong> {item.get('text', '')}<br>
                        <em>{comments}</em>
                    </div>
'''
        
        html += '''
                </div>
            </div>
            
            <!-- Footer -->
            <div class="footer">
                <div class="footer-logo">
                    <div class="bar"></div>
                    <div class="bar"></div>
                </div>
                <div class="footer-text">
                    Generated by RedlineAI &bull; abonmarche.com
                </div>
            </div>
        </div>
    </div>
</body>
</html>
'''
        
        return html
    
    def generate_word_report(self, review_results: Dict[str, Any]) -> io.BytesIO:
        """
        Generate a Word document report matching Abonmarche checklist format.
        """
        doc = Document()
        project = review_results.get('project_summary', {})
        summary = review_results.get('summary', {})
        phase = review_results.get('phase', '')
        
        # Get form metadata
        form_project_name = review_results.get('form_project_name') or project.get('project_name', '')
        form_project_number = review_results.get('form_project_number') or project.get('project_number', '')
        form_project_manager = review_results.get('form_project_manager', '')
        form_reviewer = review_results.get('form_reviewer', 'RedlineAI')
        
        # Set up styles
        style = doc.styles['Normal']
        style.font.name = 'Arial'
        style.font.size = Pt(10)
        
        # Header
        header = doc.add_heading(f'{phase} ENGINEERING QA/QC Review', level=1)
        header.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Project info table
        info_table = doc.add_table(rows=5, cols=2)
        info_table.style = 'Table Grid'
        
        info_data = [
            ('Project No.:', form_project_number),
            ('Project Name:', form_project_name),
            ('Project Manager:', form_project_manager),
            ('Reviewer:', form_reviewer),
            ('Date:', datetime.now().strftime('%m/%d/%Y'))
        ]
        
        for i, (label, value) in enumerate(info_data):
            row = info_table.rows[i]
            row.cells[0].text = label
            row.cells[0].paragraphs[0].runs[0].bold = True
            row.cells[1].text = str(value) if value else ''
        
        doc.add_paragraph()
        
        # Summary
        summary_para = doc.add_paragraph()
        summary_para.add_run('SUMMARY: ').bold = True
        summary_para.add_run(
            f"Total: {summary.get('total_items', 0)} | "
            f"Passed: {summary.get('yes_count', 0)} | "
            f"Failed: {summary.get('no_count', 0)} | "
            f"N/A: {summary.get('na_count', 0)}"
        )
        
        doc.add_paragraph()
        
        # Checklist sections
        for section in review_results.get('sections', []):
            # Section header
            section_header = doc.add_paragraph()
            run = section_header.add_run(section['title'])
            run.bold = True
            run.underline = True
            
            # Items table
            table = doc.add_table(rows=1, cols=5)
            table.style = 'Table Grid'
            
            # Header row
            header_row = table.rows[0]
            headers = ['', 'YES', 'NO', 'N/A', 'COMMENTS']
            for i, header_text in enumerate(headers):
                cell = header_row.cells[i]
                cell.text = header_text
                if i > 0:
                    cell.paragraphs[0].runs[0].bold = True
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Item rows
            for item in section.get('items', []):
                row = table.add_row()
                row.cells[0].text = item.get('text', '')
                
                status = item.get('status', 'N/A')
                
                # Mark the appropriate checkbox
                row.cells[1].text = '✓' if status == 'YES' else '☐'
                row.cells[2].text = '✓' if status == 'NO' else '☐'
                row.cells[3].text = '✓' if status == 'N/A' else '☐'
                
                for i in range(1, 4):
                    row.cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                comments = item.get('comments', '')
                if len(comments) > 50:
                    comments = comments[:50] + '...'
                row.cells[4].text = comments
            
            doc.add_paragraph()
        
        # Items for next phase
        items_for_next = summary.get('items_for_next_phase', [])
        next_phase = review_results.get('next_phase')
        
        if items_for_next and next_phase:
            doc.add_heading(f'Items Required to Advance to {next_phase}', level=2)
            
            for item in items_for_next:
                para = doc.add_paragraph(style='List Bullet')
                para.add_run(f"{item.get('id', '')}: ").bold = True
                para.add_run(item.get('text', ''))
        
        # Review Comments
        doc.add_heading('Review Comments', level=2)
        
        for section in review_results.get('sections', []):
            for item in section.get('items', []):
                if item.get('status') == 'NO':
                    para = doc.add_paragraph()
                    para.add_run(f"{item.get('id', '')}: ").bold = True
                    para.add_run(item.get('text', ''))
        
        # Save to buffer
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        return buffer
    
    def generate_excel_quantities(self, review_results: Dict[str, Any]) -> io.BytesIO:
        """
        Generate an Excel spreadsheet with extracted quantities.
        """
        import openpyxl
        from openpyxl.styles import Font, PatternFill, Alignment
        
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = "Quantities"
        
        # Header styling
        header_font = Font(bold=True, color="FFFFFF")
        header_fill = PatternFill(start_color="1a365d", end_color="1a365d", fill_type="solid")
        
        # Headers
        headers = ['Sheet', 'Item', 'Quantity', 'Unit', 'Location']
        for col, header in enumerate(headers, 1):
            cell = ws.cell(row=1, column=col, value=header)
            cell.font = header_font
            cell.fill = header_fill
        
        # Data
        row_num = 2
        for qty_group in review_results.get('quantities', []):
            sheet_name = qty_group.get('sheet', '')
            for qty in qty_group.get('quantities', []):
                ws.cell(row=row_num, column=1, value=sheet_name)
                ws.cell(row=row_num, column=2, value=qty.get('item', ''))
                ws.cell(row=row_num, column=3, value=qty.get('quantity', ''))
                ws.cell(row=row_num, column=4, value=qty.get('unit', ''))
                ws.cell(row=row_num, column=5, value=qty.get('location', ''))
                row_num += 1
        
        # Auto-size columns
        for col in ws.columns:
            max_length = 0
            column = col[0].column_letter
            for cell in col:
                try:
                    if len(str(cell.value)) > max_length:
                        max_length = len(str(cell.value))
                except:
                    pass
            adjusted_width = min(max_length + 2, 50)
            ws.column_dimensions[column].width = adjusted_width
        
        buffer = io.BytesIO()
        wb.save(buffer)
        buffer.seek(0)
        
        return buffer
